"""Check the structure of the 2024 DOCX file"""
from docx import Document

doc = Document(r"C:\2026_AI_Collaboration\aiseries\data\2024_AI_Programme.docx")

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

print("\n=== First 20 paragraphs ===")
for i, para in enumerate(doc.paragraphs[:20]):
    if para.text.strip():
        print(f"{i}: {para.text[:100]}")

print("\n=== Tables structure ===")
for i, table in enumerate(doc.tables):
    print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    if len(table.rows) > 0:
        print("  Headers:", [cell.text.strip()[:30] for cell in table.rows[0].cells])
        if len(table.rows) > 1:
            print("  Row 1:", [cell.text.strip()[:30] for cell in table.rows[1].cells])
